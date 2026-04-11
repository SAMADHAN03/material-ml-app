import streamlit as st
import joblib
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import io
import tempfile

from mp_api.client import MPRester
from sklearn.metrics import mean_absolute_error, r2_score

from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, Image
from reportlab.lib.styles import getSampleStyleSheet

from docx import Document

# -----------------------------
# CONFIG
# -----------------------------
st.set_page_config(page_title="Material ML Analysis", layout="wide")
API_KEY = st.secrets["MP_API_KEY"]

# -----------------------------
# LOAD MODEL
# -----------------------------
@st.cache_resource
def load_assets():
    model = joblib.load("model.pkl")
    features = joblib.load("features.pkl")
    return model, features

model, feature_columns = load_assets()

# -----------------------------
# API FUNCTION
# -----------------------------
@st.cache_data
def get_base_gap(material):
    try:
        with MPRester(API_KEY) as mpr:
            docs = mpr.summary.search(formula=material, fields=["band_gap"])
            if docs:
                return docs[0].band_gap
    except:
        pass
    return 3.0

# -----------------------------
# VARSHNI MODEL
# -----------------------------
def varshni(material, T):
    params = {
        "ZnO": (5.5e-4, 900),
        "Fe2O3": (4.5e-4, 500),
        "CeO2": (4.7e-4, 600)
    }
    alpha, beta = params.get(material, (5e-4, 500))
    Eg0 = get_base_gap(material)
    return Eg0 - (alpha * T**2) / (T + beta)

# -----------------------------
# ML PIPELINE
# -----------------------------
def run_pipeline(material, dopant, temp, conc, size):

    mp_gap = get_base_gap(material)
    oqmd_gap = 3.10
    aflow_gap = 3.05
    theoretical = varshni(material, temp)

    df_input = pd.DataFrame([{
        "material": material,
        "dopant": dopant,
        "temp": temp,
        "conc": conc,
        "particle_size": size
    }])

    df_encoded = pd.get_dummies(df_input)
    df_encoded = df_encoded.reindex(columns=feature_columns, fill_value=0)

    predicted = model.predict(df_encoded)[0]
    expected = np.mean([predicted, theoretical, oqmd_gap, aflow_gap])

    return predicted, expected, theoretical, mp_gap

# -----------------------------
# PDF GENERATOR
# -----------------------------
def generate_pdf(df_results, fig, mae=None, r2=None):

    file_path = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf").name

    doc = SimpleDocTemplate(file_path)
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph("Material ML Analysis Report", styles['Title']))
    elements.append(Spacer(1, 12))

    if mae is not None and r2 is not None:
        elements.append(Paragraph(f"MAE: {mae:.4f}", styles['Normal']))
        elements.append(Paragraph(f"R² Score: {r2:.4f}", styles['Normal']))
        elements.append(Spacer(1, 12))

    data = [df_results.columns.tolist()] + df_results.values.tolist()
    elements.append(Table(data))
    elements.append(Spacer(1, 12))

    img_path = tempfile.NamedTemporaryFile(delete=False, suffix=".png").name
    fig.savefig(img_path)

    elements.append(Paragraph("Band Gap Comparison Graph", styles['Heading2']))
    elements.append(Image(img_path, width=400, height=300))

    doc.build(elements)

    return file_path

# -----------------------------
# WORD GENERATOR
# -----------------------------
def generate_word(df_results, mae=None, r2=None):

    file_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name

    doc = Document()
    doc.add_heading("Material ML Analysis Report", 0)

    if mae is not None and r2 is not None:
        doc.add_paragraph(f"MAE: {mae:.4f}")
        doc.add_paragraph(f"R² Score: {r2:.4f}")

    table = doc.add_table(rows=1, cols=len(df_results.columns))

    for i, col in enumerate(df_results.columns):
        table.rows[0].cells[i].text = col

    for _, row in df_results.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

    doc.save(file_path)

    return file_path

# -----------------------------
# UI
# -----------------------------
st.title("🔬 Advanced Material ML Analysis")

uploaded_file = st.file_uploader("📂 Upload CSV or Excel", type=["csv", "xlsx"])

if uploaded_file:

    # READ FILE
    if uploaded_file.name.endswith(".csv"):
        df = pd.read_csv(uploaded_file, sep=None, engine="python")
    else:
        df = pd.read_excel(uploaded_file)

    st.subheader("📊 Raw Data")
    st.dataframe(df)

    # -----------------------------
    # CLEANING
    # -----------------------------
    before_rows = len(df)

    df.columns = df.columns.str.strip().str.lower()
    has_exp = "band_gap_exp" in df.columns

    for col in ["temp", "conc", "particle_size", "band_gap_exp"]:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce")

    df = df[df["temp"] > 0]

    df["conc"] = df.get("conc", 0).fillna(0)
    df["particle_size"] = df.get("particle_size", 0).fillna(0)

    df = df.dropna(subset=["material", "temp"])
    df = df.reset_index(drop=True)

    after_rows = len(df)

    st.subheader("🧹 Cleaning Report")
    st.write(f"Before: {before_rows} | After: {after_rows}")

    st.subheader("✅ Cleaned Data")
    st.dataframe(df)

    # -----------------------------
    # RUN ANALYSIS
    # -----------------------------
    if st.button("🚀 Run Analysis"):

        with st.spinner("Running ML Analysis..."):

            results = []
            progress = st.progress(0.0)

            for i, row in df.iterrows():

                pred, exp, theo, mp = run_pipeline(
                    row["material"],
                    row.get("dopant", "None"),
                    row["temp"],
                    row.get("conc", 0),
                    row.get("particle_size", 0)
                )

                res = {
                    "Material": row["material"],
                    "Predicted": pred,
                    "Expected": exp,
                    "Theoretical": theo,
                    "MP": mp
                }

                if has_exp:
                    res["Experimental"] = row.get("band_gap_exp", np.nan)

                results.append(res)
                progress.progress((i + 1) / len(df))

        df_results = pd.DataFrame(results)

        st.subheader("📊 Results")
        st.dataframe(df_results)

        # -----------------------------
        # METRICS
        # -----------------------------
        mae, r2 = None, None

        if has_exp:
            df_valid = df_results.dropna(subset=["Experimental"])

            if not df_valid.empty:
                mae = mean_absolute_error(df_valid["Experimental"], df_valid["Predicted"])
                r2 = r2_score(df_valid["Experimental"], df_valid["Predicted"])

                col1, col2 = st.columns(2)
                col1.metric("MAE", round(mae, 4))
                col2.metric("R²", round(r2, 4))

        # -----------------------------
        # GRAPH
        # -----------------------------
        fig, ax = plt.subplots()

        ax.plot(df_results["Predicted"], label="Predicted")
        ax.plot(df_results["Expected"], label="Expected")

        if has_exp:
            ax.plot(df_results["Experimental"], label="Experimental")

        ax.legend()
        ax.set_title("Band Gap Comparison")

        st.pyplot(fig)

        # -----------------------------
        # DOWNLOAD CSV
        # -----------------------------
        csv = df_results.to_csv(index=False).encode()
        st.download_button("⬇ Download CSV", csv, "results.csv")

        # -----------------------------
        # DOWNLOAD GRAPH
        # -----------------------------
        buf = io.BytesIO()
        fig.savefig(buf, format="png")
        st.download_button("⬇ Download Graph", buf.getvalue(), "graph.png")

        # -----------------------------
        # PDF DOWNLOAD
        # -----------------------------
        pdf_file = generate_pdf(df_results, fig, mae, r2)

        with open(pdf_file, "rb") as f:
            st.download_button("📄 Download PDF", f, "report.pdf")

        # -----------------------------
        # WORD DOWNLOAD
        # -----------------------------
        word_file = generate_word(df_results, mae, r2)

        with open(word_file, "rb") as f:
            st.download_button("📝 Download Word", f, "report.docx")

        st.success("✅ Analysis Completed Successfully!")
